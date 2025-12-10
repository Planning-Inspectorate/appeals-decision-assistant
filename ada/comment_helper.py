"""Helper class for adding comments to PDF and Word documents."""

import logging
import re
from pathlib import Path

from docx import Document
from spire.pdf import *
from spire.pdf.common import *
from spire.pdf.annotations.PdfTextMarkupAnnotation import PdfTextMarkupAnnotation
from spire.pdf import PdfTextExtractor, PdfTextExtractOptions


class CommentHelper:
    """Helper for adding LLM-generated comments to PDF and Word documents."""

    @staticmethod
    def extract_line_numbers(text: str) -> list[int]:
        """Extract line numbers from comment text like 'Lines 9-11' or 'line 53'."""
        line_numbers = []
        # Match patterns like "Lines 9-11", "line 53", "Lines 106-108"
        patterns = [
            r'[Ll]ines?\s+(\d+)-(\d+)',  # Lines 9-11
            r'[Ll]ines?\s+(\d+)',          # line 53
        ]

        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                if len(match.groups()) == 2:
                    start, end = int(match.group(1)), int(match.group(2))
                    line_numbers.extend(range(start, end + 1))
                else:
                    line_numbers.append(int(match.group(1)))

        return sorted(set(line_numbers))

    @staticmethod
    def _parse_comment_sections(comments_text: str) -> list[str]:
        """Parse comments text into individual comment sections."""
        comment_sections = []
        current_section = []

        for line in comments_text.split("\n"):
            line = line.strip()
            if line and (line[0].isdigit() or line.startswith("Location:")):
                if current_section:
                    comment_sections.append("\n".join(current_section))
                current_section = [line]
            elif current_section:
                current_section.append(line)

        if current_section:
            comment_sections.append("\n".join(current_section))

        return comment_sections

    @staticmethod
    def add_comments_to_pdf(pdf_path: Path, comments_text: str, output_path: Path) -> None:
        """Add LLM improvement suggestions as highlight annotations on the PDF."""
        doc = PdfDocument()
        doc.LoadFromFile(str(pdf_path))

        all_text = ""
        for page_idx in range(doc.Pages.Count):
            page = doc.Pages[page_idx]
            extractor = PdfTextExtractor(page)
            options = PdfTextExtractOptions()
            all_text += extractor.ExtractText(options) + "\n"

        text_lines = all_text.split("\n")

        comment_sections = CommentHelper._parse_comment_sections(comments_text)

        annotations_added = 0
        for comment in comment_sections[:50]:
            line_numbers = CommentHelper.extract_line_numbers(comment)

            if not line_numbers:
                continue

            target_lines = []
            for line_num in line_numbers:
                if 1 <= line_num <= len(text_lines):
                    target_lines.append(text_lines[line_num - 1])

            if not target_lines:
                continue

            search_text = target_lines[0].strip()
            if not search_text or len(search_text) < 3:
                continue

            for page_idx in range(doc.Pages.Count):
                page = doc.Pages[page_idx]

                finder = PdfTextFinder(page)
                finder.Options.Parameter = TextFindParameter.IgnoreCase

                fragments = finder.Find(search_text)

                if fragments and len(fragments) > 0:
                    text_fragment = fragments[0]

                    for i in range(len(text_fragment.Bounds)):
                        rect = text_fragment.Bounds[i]

                        annotation = PdfTextMarkupAnnotation("ADA", comment[:500] if len(comment) > 500 else comment, rect)
                        annotation.TextMarkupAnnotationType = PdfTextMarkupAnnotationType.Highlight
                        annotation.TextMarkupColor = PdfRGBColor(Color.get_Yellow())

                        page.AnnotationsWidget.Add(annotation)

                    annotations_added += 1
                    break

        doc.SaveToFile(str(output_path))
        doc.Close()

        logging.info(f"Added {annotations_added} highlight annotations to PDF")

    @staticmethod
    def add_comments_to_docx(docx_path: Path, comments_text: str, output_path: Path) -> None:
        """Add LLM improvement suggestions as comments on the Word document."""
        doc = Document(str(docx_path))

        comment_sections = CommentHelper._parse_comment_sections(comments_text)

        comments_added = 0
        for comment in comment_sections[:50]:  # Limit to 50 comments
            line_numbers = CommentHelper.extract_line_numbers(comment)

            if not line_numbers:
                continue

            for line_num in line_numbers:
                para_idx = line_num - 1

                if 0 <= para_idx < len(doc.paragraphs):
                    paragraph = doc.paragraphs[para_idx]

                    if paragraph.runs:
                        doc.add_comment(
                            runs=paragraph.runs,
                            text=comment[:500] if len(comment) > 500 else comment,
                            author="ADA",
                            initials="ADA",
                        )
                        comments_added += 1
                        break

        doc.save(str(output_path))

        logging.info(f"Added {comments_added} comments to Word document")
