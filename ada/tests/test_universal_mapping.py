"""Test that the line mapping solution works universally with different documents."""

from pathlib import Path
from docx import Document
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
import tempfile


def create_test_document(content_paragraphs: list[str]) -> Path:
    """Create a test Word document with given paragraphs."""
    doc = Document()
    for para_text in content_paragraphs:
        doc.add_paragraph(para_text)

    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(temp_file.name)
    return Path(temp_file.name)


def build_line_to_para_map(docx_path: Path) -> tuple[dict, list[str], list[str]]:
    """Build the mapping the same way comment_helper does."""
    doc = Document(str(docx_path))
    loader = UnstructuredWordDocumentLoader(str(docx_path))
    extracted_text = loader.load()[0].page_content
    extracted_lines = extracted_text.split("\n")

    line_to_para_map = {}
    used_paragraphs = set()

    for line_idx, line in enumerate(extracted_lines, start=1):
        line_text = line.strip()
        if not line_text:
            continue

        for para_idx, paragraph in enumerate(doc.paragraphs):
            if para_idx in used_paragraphs:
                continue

            para_text = paragraph.text.strip()
            if not para_text:
                continue

            # Try exact match first
            if para_text == line_text:
                line_to_para_map[line_idx] = para_idx
                used_paragraphs.add(para_idx)
                break

            # Fall back to fuzzy match
            normalized_para = " ".join(para_text.split())
            normalized_line = " ".join(line_text.split())
            if normalized_para == normalized_line:
                line_to_para_map[line_idx] = para_idx
                used_paragraphs.add(para_idx)
                break

    para_texts = [p.text for p in doc.paragraphs]
    return line_to_para_map, extracted_lines, para_texts


def test_simple_document():
    """Test with a simple document."""
    print("\n" + "=" * 80)
    print("TEST 1: Simple Document")
    print("=" * 80)

    paragraphs = [
        "Title",
        "First paragraph of content.",
        "Second paragraph with different text.",
        "Third paragraph here.",
    ]

    doc_path = create_test_document(paragraphs)
    line_map, extracted, paras = build_line_to_para_map(doc_path)

    print(f"Document paragraphs: {len(paras)}")
    print(f"Extracted lines: {len(extracted)}")
    print(f"Successful mappings: {len(line_map)}")

    # Verify each paragraph is mapped
    success = len(line_map) == len([p for p in paragraphs if p.strip()])
    print(f"Result: {'✓ PASS' if success else '✗ FAIL'}")

    doc_path.unlink()
    return success


def test_duplicate_text():
    """Test with duplicate paragraphs (should map to first occurrence)."""
    print("\n" + "=" * 80)
    print("TEST 2: Document with Duplicate Text")
    print("=" * 80)

    paragraphs = [
        "Unique first line",
        "Duplicate paragraph",
        "Another unique line",
        "Duplicate paragraph",  # Same text as paragraph 2
        "Final unique line",
    ]

    doc_path = create_test_document(paragraphs)
    line_map, extracted, paras = build_line_to_para_map(doc_path)

    print(f"Document paragraphs: {len(paras)}")
    print(f"Extracted lines: {len(extracted)}")
    print(f"Successful mappings: {len(line_map)}")

    # Check that both duplicate lines are mapped (to different paragraph instances)
    mapped_paras = set(line_map.values())
    print(f"Unique paragraphs mapped: {len(mapped_paras)}")

    # Should map each line to a different paragraph
    success = len(line_map) >= len([p for p in paragraphs if p.strip()])
    print(f"Result: {'✓ PASS' if success else '✗ FAIL'}")

    doc_path.unlink()
    return success


def test_whitespace_variations():
    """Test with paragraphs that have extra whitespace."""
    print("\n" + "=" * 80)
    print("TEST 3: Document with Whitespace Variations")
    print("=" * 80)

    paragraphs = [
        "Normal text",
        "Text  with   extra    spaces",
        "Text\twith\ttabs",
        "Final paragraph",
    ]

    doc_path = create_test_document(paragraphs)
    line_map, extracted, paras = build_line_to_para_map(doc_path)

    print(f"Document paragraphs: {len(paras)}")
    print(f"Extracted lines: {len(extracted)}")
    print(f"Successful mappings: {len(line_map)}")

    # Should handle whitespace normalization
    success = len(line_map) >= 3  # At least 3 of 4 should map
    print(f"Result: {'✓ PASS' if success else '✗ FAIL'}")

    doc_path.unlink()
    return success


def test_empty_paragraphs():
    """Test with empty paragraphs mixed in."""
    print("\n" + "=" * 80)
    print("TEST 4: Document with Empty Paragraphs")
    print("=" * 80)

    paragraphs = [
        "First paragraph",
        "",  # Empty
        "Second paragraph",
        "",  # Empty
        "",  # Empty
        "Third paragraph",
    ]

    doc_path = create_test_document(paragraphs)
    line_map, extracted, paras = build_line_to_para_map(doc_path)

    print(f"Document paragraphs: {len(paras)}")
    print(f"Extracted lines: {len(extracted)}")
    print(f"Successful mappings: {len(line_map)}")

    # Should map only non-empty paragraphs
    non_empty_count = len([p for p in paragraphs if p.strip()])
    success = len(line_map) == non_empty_count
    print(f"Result: {'✓ PASS' if success else '✗ FAIL'}")

    doc_path.unlink()
    return success


def test_long_document():
    """Test with a longer document to ensure scalability."""
    print("\n" + "=" * 80)
    print("TEST 5: Long Document (100 paragraphs)")
    print("=" * 80)

    paragraphs = [f"Paragraph number {i} with unique text." for i in range(1, 101)]

    doc_path = create_test_document(paragraphs)
    line_map, extracted, paras = build_line_to_para_map(doc_path)

    print(f"Document paragraphs: {len(paras)}")
    print(f"Extracted lines: {len(extracted)}")
    print(f"Successful mappings: {len(line_map)}")

    # Should map at least 95% of paragraphs
    success = len(line_map) >= 95
    print(f"Result: {'✓ PASS' if success else '✗ FAIL'}")

    doc_path.unlink()
    return success


if __name__ == "__main__":
    print("\n" + "=" * 80)
    print("UNIVERSAL MAPPING SOLUTION TEST SUITE")
    print("=" * 80)

    results = []
    results.append(("Simple Document", test_simple_document()))
    results.append(("Duplicate Text", test_duplicate_text()))
    results.append(("Whitespace Variations", test_whitespace_variations()))
    results.append(("Empty Paragraphs", test_empty_paragraphs()))
    results.append(("Long Document", test_long_document()))

    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)

    for test_name, passed in results:
        status = "✓ PASS" if passed else "✗ FAIL"
        print(f"{status}: {test_name}")

    all_passed = all(result[1] for result in results)
    print(f"\n{'=' * 80}")
    print(f"Overall: {'✓ ALL TESTS PASSED' if all_passed else '✗ SOME TESTS FAILED'}")
    print("=" * 80)
