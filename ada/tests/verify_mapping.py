"""Verify that the line-to-paragraph mapping is working correctly."""

from pathlib import Path
from docx import Document
from langchain_community.document_loaders import UnstructuredWordDocumentLoader

docx_path = Path.home() / "Downloads" / "Appeal-Decision=6000640.docx"
doc = Document(str(docx_path))

# Load the document the same way the LLM saw it
loader = UnstructuredWordDocumentLoader(str(docx_path))
extracted_text = loader.load()[0].page_content
extracted_lines = extracted_text.split("\n")

# Build the mapping
line_to_para_map = {}
for line_idx, line in enumerate(extracted_lines, start=1):
    line_text = line.strip()
    if not line_text:
        continue

    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text.strip()
        if para_text and para_text == line_text:
            line_to_para_map[line_idx] = para_idx
            break

print("=" * 80)
print("VERIFICATION: LLM Line Numbers → Document Paragraphs")
print("=" * 80)

# Test some examples from the LLM output
test_cases = [
    (3, "ite visit made n 14 October"),
    (5, "025"),
    (7, "y R Dickson Sc (Hons)"),
    (11, "n Inspector appointed y the Secretary of tate"),
    (47, "section ofdropped kerb"),
    (89, "An appeal at concerning 37 Woodmill Lane"),
]

for line_num, expected_text in test_cases:
    para_idx = line_to_para_map.get(line_num)
    if para_idx is not None:
        actual_text = doc.paragraphs[para_idx].text[:50]
        extracted_line = extracted_lines[line_num - 1][:50] if line_num <= len(extracted_lines) else "N/A"
        status = "✓" if expected_text in actual_text else "✗"
        print(f"\n{status} Line {line_num} → Paragraph {para_idx}")
        print(f"  Extracted: {extracted_line}")
        print(f"  Document:  {actual_text}")
    else:
        print(f"\n✗ Line {line_num} → NOT MAPPED")

print(f"\n{'=' * 80}")
print(f"Total mappings created: {len(line_to_para_map)}")
print(f"Total extracted lines: {len(extracted_lines)}")
print(f"Total document paragraphs: {len(doc.paragraphs)}")
